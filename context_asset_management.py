"""CRUD API for org-level context assets (knowledge base)."""

from __future__ import annotations

import asyncio
import io
import concurrent.futures
import logging
import uuid
from datetime import datetime, timezone
from typing import Any

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from pydantic import BaseModel, Field

from auth_dependency import require_org_member, AuthenticatedUser
from supabase_client import supabase

logger = logging.getLogger(__name__)

router = APIRouter(tags=["context-assets"])

_ALLOWED_TYPES = {"text", "document", "url", "notion_page"}
_MAX_CONTENT_CHARS = 500_000


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def _extract_text_from_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def _extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs).strip()


def _extract_text_from_url(url: str) -> str:
    import trafilatura
    downloaded = trafilatura.fetch_url(url)
    if not downloaded:
        raise ValueError(f"Could not fetch URL: {url}")
    text = trafilatura.extract(downloaded)
    if not text:
        raise ValueError(f"Could not extract text from URL: {url}")
    return text.strip()


_thread_pool = concurrent.futures.ThreadPoolExecutor(max_workers=2)


def _bg_index_asset(asset_id: str, org_id: str, content: str, metadata: dict):
    """Fire-and-forget: chunk + embed asset, update embedding_status in metadata."""
    from context_embeddings import index_asset
    success = index_asset(asset_id, org_id, content)
    _status = "indexed" if success else "failed"
    try:
        supabase.table("context_assets").update(
            {"metadata": {**metadata, "embedding_status": _status}}
        ).eq("id", asset_id).execute()
    except Exception:
        logger.exception("Failed to update embedding_status for asset %s", asset_id)


# ---------------------------------------------------------------------------
# Request models
# ---------------------------------------------------------------------------

class CreateContextAssetRequest(BaseModel):
    org_id: str
    name: str = Field(..., min_length=1, max_length=255)
    description: str | None = None
    asset_type: str = "text"
    content: str | None = None
    source_ref: dict[str, Any] | None = None


class UpdateContextAssetRequest(BaseModel):
    name: str | None = Field(None, min_length=1, max_length=255)
    description: str | None = None
    content: str | None = None
    source_ref: dict[str, Any] | None = None


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

# Upload must be defined BEFORE {org_id} routes to avoid path matching conflicts
@router.post("/context-assets/upload")
async def upload_context_asset(
    file: UploadFile = File(...),
    name: str = Form(...),
    org_id: str = Form(...),
    description: str = Form(""),
    auth_user: AuthenticatedUser = Depends(require_org_member),
):
    file_bytes = await file.read()
    if len(file_bytes) > 10 * 1024 * 1024:
        raise HTTPException(400, "File exceeds 10MB limit")
    filename = file.filename or "unknown"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        content = _extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        content = _extract_text_from_docx(file_bytes)
    elif ext in ("txt", "md", "text", "markdown"):
        content = file_bytes.decode("utf-8", errors="replace")
    else:
        raise HTTPException(400, f"Unsupported file type: .{ext}. Supported: .pdf, .docx, .txt, .md")
    if len(content) > _MAX_CONTENT_CHARS:
        raise HTTPException(400, f"Extracted text exceeds {_MAX_CONTENT_CHARS} character limit")
    metadata = {"original_filename": filename, "file_type": ext, "char_count": len(content), "word_count": len(content.split()), "embedding_status": "pending"}
    row = {"id": str(uuid.uuid4()), "org_id": org_id, "name": name, "description": description or None, "asset_type": "document", "content": content, "metadata": metadata, "status": "active"}
    result = supabase.table("context_assets").insert(row).execute()
    asset = result.data[0] if result.data else row
    _thread_pool.submit(_bg_index_asset, asset["id"], org_id, content, metadata)
    return asset


@router.get("/context-assets/{org_id}")
async def list_context_assets(
    org_id: str,
    auth_user: AuthenticatedUser = Depends(require_org_member),
    asset_type: str | None = None,
):
    query = (
        supabase.table("context_assets")
        .select("id, org_id, name, description, asset_type, metadata, status, created_at, updated_at")
        .eq("org_id", org_id)
        .neq("status", "deleted")
        .order("created_at", desc=True)
    )
    if asset_type:
        query = query.eq("asset_type", asset_type)
    result = query.execute()
    return result.data or []


@router.post("/context-assets")
async def create_context_asset(
    body: CreateContextAssetRequest,
    auth_user: AuthenticatedUser = Depends(require_org_member),
):
    if body.asset_type not in _ALLOWED_TYPES:
        raise HTTPException(400, f"asset_type must be one of: {', '.join(sorted(_ALLOWED_TYPES))}")

    org_id = body.org_id
    status = "active"
    content = body.content
    metadata: dict[str, Any] = {}

    # --- Text assets ---
    if body.asset_type == "text":
        if not content:
            raise HTTPException(400, "content is required for text assets")
        if len(content) > _MAX_CONTENT_CHARS:
            raise HTTPException(400, f"content exceeds maximum of {_MAX_CONTENT_CHARS} characters")

    # --- URL assets ---
    elif body.asset_type == "url":
        if not body.source_ref or not body.source_ref.get("url"):
            raise HTTPException(400, "source_ref.url is required for url assets")
        url = body.source_ref["url"]
        try:
            content = _extract_text_from_url(url)
        except ValueError as exc:
            raise HTTPException(422, str(exc))
        if len(content) > _MAX_CONTENT_CHARS:
            raise HTTPException(400, f"Extracted text exceeds {_MAX_CONTENT_CHARS} character limit")
        metadata["source_url"] = url
        metadata["embedding_status"] = "pending"

    # --- Notion page assets ---
    elif body.asset_type == "notion_page":
        if not body.source_ref or not body.source_ref.get("notion_page_id"):
            raise HTTPException(400, "source_ref.notion_page_id is required for notion_page assets")
        page_id = body.source_ref["notion_page_id"]
        # Retrieve Notion API key from org_secrets
        secret_row = (
            supabase.table("org_secrets")
            .select("encrypted_value")
            .eq("org_id", org_id)
            .eq("key", "NOTION_API_KEY")
            .execute()
        )
        if not secret_row.data:
            raise HTTPException(400, "Notion API key not configured for this organization")
        from encryption_utils import decrypt_value
        notion_token = decrypt_value(secret_row.data[0]["encrypted_value"])
        from context_notion import fetch_notion_page
        try:
            content = fetch_notion_page(page_id, notion_token)
        except Exception as exc:
            raise HTTPException(422, f"Failed to fetch Notion page: {exc}")
        if len(content) > _MAX_CONTENT_CHARS:
            raise HTTPException(400, f"Extracted text exceeds {_MAX_CONTENT_CHARS} character limit")
        metadata["notion_page_id"] = page_id
        metadata["embedding_status"] = "pending"

    # --- Document (handled by upload endpoint) ---
    else:
        status = "processing"

    if content:
        metadata["char_count"] = len(content)
        metadata["word_count"] = len(content.split())
    if body.asset_type == "text":
        metadata["embedding_status"] = "pending"

    row = {
        "id": str(uuid.uuid4()),
        "org_id": org_id,
        "name": body.name,
        "description": body.description,
        "asset_type": body.asset_type,
        "content": content,
        "source_ref": body.source_ref,
        "metadata": metadata,
        "status": status,
    }
    result = supabase.table("context_assets").insert(row).execute()
    asset = result.data[0] if result.data else row

    # Fire-and-forget embedding indexing for assets with content
    if content and body.asset_type in ("text", "url", "notion_page"):
        _thread_pool.submit(_bg_index_asset, asset["id"], org_id, content, metadata)

    return asset


@router.get("/context-assets/{org_id}/{asset_id}")
async def get_context_asset(
    org_id: str,
    asset_id: str,
    auth_user: AuthenticatedUser = Depends(require_org_member),
):
    result = (
        supabase.table("context_assets")
        .select("*")
        .eq("id", asset_id)
        .eq("org_id", org_id)
        .neq("status", "deleted")
        .execute()
    )
    if not result.data:
        raise HTTPException(404, "Context asset not found")
    return result.data[0]


@router.put("/context-assets/{org_id}/{asset_id}")
async def update_context_asset(
    org_id: str,
    asset_id: str,
    body: UpdateContextAssetRequest,
    auth_user: AuthenticatedUser = Depends(require_org_member),
):
    existing = (
        supabase.table("context_assets")
        .select("id, asset_type")
        .eq("id", asset_id)
        .eq("org_id", org_id)
        .neq("status", "deleted")
        .execute()
    )
    if not existing.data:
        raise HTTPException(404, "Context asset not found")

    updates: dict[str, Any] = {"updated_at": datetime.now(timezone.utc).isoformat()}

    if body.name is not None:
        updates["name"] = body.name
    if body.description is not None:
        updates["description"] = body.description
    if body.content is not None:
        if len(body.content) > _MAX_CONTENT_CHARS:
            raise HTTPException(400, f"content exceeds maximum of {_MAX_CONTENT_CHARS} characters")
        updates["content"] = body.content
        updates["metadata"] = {
            "char_count": len(body.content),
            "word_count": len(body.content.split()),
        }
    if body.source_ref is not None:
        updates["source_ref"] = body.source_ref

    result = (
        supabase.table("context_assets")
        .update(updates)
        .eq("id", asset_id)
        .eq("org_id", org_id)
        .execute()
    )
    return result.data[0] if result.data else updates


@router.post("/context-assets/{org_id}/{asset_id}/sync")
async def sync_context_asset(
    org_id: str,
    asset_id: str,
    auth_user: AuthenticatedUser = Depends(require_org_member),
):
    existing = (
        supabase.table("context_assets")
        .select("*")
        .eq("id", asset_id)
        .eq("org_id", org_id)
        .neq("status", "deleted")
        .execute()
    )
    if not existing.data:
        raise HTTPException(404, "Context asset not found")

    asset = existing.data[0]
    asset_type = asset["asset_type"]
    source_ref = asset.get("source_ref") or {}
    metadata = asset.get("metadata") or {}

    if asset_type == "url":
        url = source_ref.get("url")
        if not url:
            raise HTTPException(400, "No source URL stored for this asset")
        try:
            content = _extract_text_from_url(url)
        except ValueError as exc:
            raise HTTPException(422, str(exc))

    elif asset_type == "notion_page":
        page_id = source_ref.get("notion_page_id")
        if not page_id:
            raise HTTPException(400, "No Notion page ID stored for this asset")
        secret_row = (
            supabase.table("org_secrets")
            .select("encrypted_value")
            .eq("org_id", org_id)
            .eq("key", "NOTION_API_KEY")
            .execute()
        )
        if not secret_row.data:
            raise HTTPException(400, "Notion API key not configured for this organization")
        from encryption_utils import decrypt_value
        notion_token = decrypt_value(secret_row.data[0]["encrypted_value"])
        from context_notion import fetch_notion_page
        try:
            content = fetch_notion_page(page_id, notion_token)
        except Exception as exc:
            raise HTTPException(422, f"Failed to fetch Notion page: {exc}")

    else:
        raise HTTPException(400, f"Sync is not supported for asset type: {asset_type}")

    if len(content) > _MAX_CONTENT_CHARS:
        raise HTTPException(400, f"Extracted text exceeds {_MAX_CONTENT_CHARS} character limit")

    metadata["char_count"] = len(content)
    metadata["word_count"] = len(content.split())
    metadata["embedding_status"] = "pending"
    metadata["last_synced_at"] = datetime.now(timezone.utc).isoformat()

    supabase.table("context_assets").update({
        "content": content,
        "metadata": metadata,
        "status": "active",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }).eq("id", asset_id).eq("org_id", org_id).execute()

    # Re-index after sync
    _thread_pool.submit(_bg_index_asset, asset_id, org_id, content, metadata)

    return {"synced": True, "char_count": len(content), "word_count": len(content.split())}


@router.delete("/context-assets/{org_id}/{asset_id}")
async def delete_context_asset(
    org_id: str,
    asset_id: str,
    auth_user: AuthenticatedUser = Depends(require_org_member),
):
    existing = (
        supabase.table("context_assets")
        .select("id")
        .eq("id", asset_id)
        .eq("org_id", org_id)
        .neq("status", "deleted")
        .execute()
    )
    if not existing.data:
        raise HTTPException(404, "Context asset not found")

    # Clean up embedding chunks before soft-deleting
    from context_embeddings import delete_asset_chunks
    delete_asset_chunks(asset_id)

    supabase.table("context_assets").update({
        "status": "deleted",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }).eq("id", asset_id).eq("org_id", org_id).execute()

    return {"deleted": True}
